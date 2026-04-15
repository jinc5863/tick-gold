#!/usr/bin/env python3
"""将 Markdown 转换为 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import sys

# 读取 Markdown 文件
with open('/Users/office01/work/tick-gold/docs/ECC_Commands_Manual.md', 'r', encoding='utf-8') as f:
    content = f.read()

# 创建 Word 文档
doc = Document()

# 添加标题
title = doc.add_heading('ECC Commands Manual', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加副标题
subtitle = doc.add_paragraph('Tick Gold - XAUUSD 量化交易系统')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 解析内容并添加到文档
lines = content.split('\n')
in_code_block = False
code_lines = []

for line in lines:
    line = line.strip()

    if line.startswith('# '):
        doc.add_heading(line[2:], 0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], 1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], 2)
    elif line.startswith('**') and line.endswith('**'):
        para = doc.add_paragraph()
        para.add_run(line.replace('**', '')).bold = True
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('|'):
        # 表格处理简化
        continue
    elif line.startswith('```'):
        in_code_block = not in_code_block
    elif in_code_block:
        code_lines.append(line)
    elif line and not line.startswith('#'):
        if code_lines:
            # 添加代码块
            code = '\n'.join(code_lines)
            para = doc.add_paragraph()
            run = para.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            code_lines = []
        if line:
            doc.add_paragraph(line)

# 保存文档
doc.save('/Users/office01/work/tick-gold/docs/ECC_Commands_Manual.docx')
print("Word 文档已生成: docs/ECC_Commands_Manual.docx")